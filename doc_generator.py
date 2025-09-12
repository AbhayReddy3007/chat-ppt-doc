from docx import Document
from docx.shared import Pt, RGBColor

def create_doc(title, points, filename="output.docx"):
    """
    Generate a Word document from title + outline.
    - Title as main heading
    - Sections with headings
    - Bulleted points for descriptions
    """

    doc = Document()

    # --- Title Page ---
    doc.add_heading(title, level=0)

    # --- Sections ---
    for idx, item in enumerate(points, start=1):
        section_title = item.get("title", f"Section {idx}")
        description = item.get("description", "")

        # Section heading
        doc.add_heading(section_title, level=1)

        # Add bullets
        if description:
            for line in description.split("\n"):
                if line.strip():
                    p = doc.add_paragraph(line.strip())
                    p.style = "List Bullet"
                    run = p.runs[0]
                    run.font.size = Pt(11)
                    run.font.color.rgb = RGBColor(40, 40, 40)  # dark gray

    # --- Save File ---
    doc.save(filename)
    return filename
